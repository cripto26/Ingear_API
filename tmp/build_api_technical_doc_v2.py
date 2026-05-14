# -*- coding: utf-8 -*-
from __future__ import annotations

import ast
import re
from zipfile import ZipFile
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
APP_DIR = ROOT / "app"
OUTPUT = ROOT / "docs" / "INFORME_TECNICO_API_v2.docx"
LOGO_PATH = ROOT / "tmp" / "v1_media" / "image2.png"

ACCENT = "004F5E"
ACCENT_DARK = "003F4A"
ACCENT_SOFT = "E8F3F5"
HEADER_FILL = "DDEFF3"
SUBTLE_FILL = "F7FAFB"
BORDER = "A8BFC5"
TEXT = "242424"


def find_v1_document() -> Path:
    downloads = Path.home() / "Downloads"
    matches = sorted(downloads.glob("INFORME T*API v.1.docx"))
    if not matches:
        raise FileNotFoundError("No se encontro el informe tecnico API v1 en Descargas.")
    return matches[0]


def ensure_logo_asset() -> None:
    if LOGO_PATH.exists():
        return
    LOGO_PATH.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(find_v1_document()) as archive:
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
        logo_name = next((name for name in media_files if name.endswith("image2.png")), media_files[0])
        LOGO_PATH.write_bytes(archive.read(logo_name))


def direct_text(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def parse_requirements() -> list[str]:
    req_path = ROOT / "requirements.txt"
    if not req_path.exists():
        return []
    return [
        line.strip()
        for line in req_path.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]


def extract_settings_fields() -> list[dict[str, str]]:
    tree = ast.parse((APP_DIR / "core" / "config.py").read_text(encoding="utf-8"))
    fields: list[dict[str, str]] = []
    for node in tree.body:
        if not isinstance(node, ast.ClassDef) or node.name != "Settings":
            continue
        for stmt in node.body:
            if isinstance(stmt, ast.AnnAssign) and isinstance(stmt.target, ast.Name):
                default = ""
                required = "Si"
                if stmt.value is not None:
                    default = ast.unparse(stmt.value)
                    required = "No"
                fields.append(
                    {
                        "name": stmt.target.id,
                        "type": ast.unparse(stmt.annotation),
                        "required": required,
                        "default": default,
                    }
                )
    return fields


def extract_router_prefixes() -> tuple[dict[str, str], dict[str, str]]:
    api_path = APP_DIR / "api" / "v1" / "api.py"
    tree = ast.parse(api_path.read_text(encoding="utf-8"))
    prefixes: dict[str, str] = {}
    tags: dict[str, str] = {}

    for node in ast.walk(tree):
        if not (
            isinstance(node, ast.Call)
            and isinstance(node.func, ast.Attribute)
            and node.func.attr == "include_router"
        ):
            continue

        module_name = None
        if node.args and isinstance(node.args[0], ast.Attribute):
            value = node.args[0].value
            if isinstance(value, ast.Name):
                module_name = value.id

        if not module_name:
            continue

        kwargs = {kw.arg: kw.value for kw in node.keywords}
        prefix_node = kwargs.get("prefix")
        if isinstance(prefix_node, ast.Constant):
            prefixes[module_name] = str(prefix_node.value)

        tags_node = kwargs.get("tags")
        if isinstance(tags_node, ast.List) and tags_node.elts:
            first = tags_node.elts[0]
            if isinstance(first, ast.Constant):
                tags[module_name] = str(first.value)

    return prefixes, tags


def _status_code_to_text(node: ast.AST | None) -> str:
    if node is None:
        return "200"
    if isinstance(node, ast.Constant):
        return str(node.value)
    if isinstance(node, ast.Attribute):
        match = re.search(r"HTTP_(\d+)_", node.attr)
        if match:
            return match.group(1)
        return node.attr
    return ast.unparse(node)


def _default_operation_summary(method: str, path: str, function_name: str) -> str:
    clean = path.strip("/")
    last = clean.split("/")[-1] if clean else "recurso"
    if method == "GET" and "{" not in last:
        return "Lista recursos con paginacion o filtros del modulo."
    if method == "GET":
        return "Consulta el detalle del recurso solicitado."
    if method == "POST" and clean.endswith("login"):
        return "Autentica credenciales, aplica protecciones de login y emite token."
    if method == "POST":
        return "Crea un recurso o ejecuta una accion del modulo."
    if method == "PUT":
        return "Actualiza parcialmente el recurso usando campos enviados."
    if method == "DELETE":
        return "Elimina la relacion o el recurso indicado."
    return function_name.replace("_", " ").capitalize()


def route_description(tag: str, method: str, path: str, function_name: str) -> str:
    exact = {
        "/api/v1/auth/login": "Inicio de sesion con usuario/cedula/email, clave, Turnstile opcional y proteccion contra fuerza bruta.",
        "/api/v1/auth/refresh": "Rota la sesion de refresh desde cookie HttpOnly y emite un nuevo access token.",
        "/api/v1/auth/logout": "Revoca la sesion de refresh actual y limpia la cookie del cliente.",
        "/api/v1/auth/me": "Devuelve el empleado autenticado, su rol inferido y permisos de vista.",
        "/api/v1/auth/change-password": "Cambia la contrasena, revoca sesiones activas y entrega una nueva sesion.",
        "/api/v1/productos/world-office/estado": "Consulta si la integracion World Office esta habilitada, configurada y conectada.",
        "/api/v1/productos/world-office/sync": "Sincroniza existencias de World Office sobre el catalogo local de productos.",
        "/api/v1/fx/rates": "Consulta TRM USD/COP y calcula EUR/COP con fuentes BanRep y BCE.",
        "/api/v1/cuentas-cobro/prefill/proyecto/{proyecto_id}": "Calcula datos sugeridos para una cuenta de cobro a partir de proyecto, oportunidad, cliente y ultima cotizacion.",
        "/api/v1/notificaciones/marcar-todas-leidas": "Marca como leidas todas las notificaciones pendientes del usuario autenticado.",
    }
    if path in exact:
        return exact[path]
    if path.endswith("/imagen"):
        return "Genera y entrega un thumbnail JPEG desde la URL de imagen configurada para el producto."
    if path.endswith("/aprobar"):
        return "Aprueba una cotizacion pendiente si el usuario es jefe directo o Gerencia."
    if path.endswith("/versiones"):
        return "Lista el historial de versiones guardado antes de cada actualizacion."
    if "/versiones/" in path:
        return "Consulta una version especifica del historial de la cotizacion."
    if path.endswith("/enviar-email"):
        return "Envia la cotizacion PDF por Gmail usando delegacion del empleado autenticado."
    if "/proyectos/" in path and any(segment in path for segment in ["/empleados", "/clientes", "/despachos"]):
        return "Administra una relacion muchos-a-muchos asociada al proyecto."
    if "/notificaciones/" in path and path.endswith("/leer"):
        return "Marca una notificacion propia como leida."
    return _default_operation_summary(method, path, function_name)


SECURITY_BY_TAG = {
    "clientes": "Bearer + permiso comercial.clientes",
    "contactos": "Bearer + permiso comercial.contactos",
    "oportunidades": "Bearer + permiso comercial.oportunidades",
    "cotizaciones": "Bearer + permiso comercial.cotizador",
    "cuentas-cobro": "Bearer + Gerencia/Logistica/Ingenieria o permiso comercial.cuentas-cobro",
    "despachos": "Bearer + rol Gerencia o Logistica",
    "empleados": "Bearer + rol Gerencia",
    "fx": "Bearer autenticado",
    "notificaciones": "Bearer autenticado",
    "paises": "Bearer autenticado",
    "productos": "Bearer + permiso comercial.productos",
    "proyectos": "Bearer + Gerencia/Logistica/Ingenieria o permiso comercial.proyectos",
    "proyectos-relaciones": "Bearer + Gerencia/Logistica/Ingenieria o permisos comercial.cotizador/oportunidades",
}


def route_security(tag: str, path: str) -> str:
    if tag == "auth":
        if path.endswith("/login"):
            return "Credenciales; Turnstile opcional; rate limit"
        if path.endswith("/refresh") or path.endswith("/logout"):
            return "Cookie HttpOnly de refresh"
        return "Bearer autenticado"
    if path.endswith("/aprobar"):
        return "Bearer + jefe directo o Gerencia"
    return SECURITY_BY_TAG.get(tag, "Bearer autenticado")


def extract_routes() -> list[dict[str, str]]:
    prefixes, tags = extract_router_prefixes()
    rows: list[dict[str, str]] = []
    for endpoint_path in sorted((APP_DIR / "api" / "v1" / "endpoints").glob("*.py")):
        module_name = endpoint_path.stem
        tree = ast.parse(endpoint_path.read_text(encoding="utf-8"))
        tag = tags.get(module_name, module_name)
        prefix = prefixes.get(module_name, "")
        for fn in [node for node in tree.body if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef))]:
            for decorator in fn.decorator_list:
                if not (
                    isinstance(decorator, ast.Call)
                    and isinstance(decorator.func, ast.Attribute)
                    and decorator.func.attr in {"get", "post", "put", "patch", "delete"}
                ):
                    continue

                method = decorator.func.attr.upper()
                route = "/"
                if decorator.args and isinstance(decorator.args[0], ast.Constant):
                    route = str(decorator.args[0].value)
                kwargs = {kw.arg: kw.value for kw in decorator.keywords}
                response_model = ast.unparse(kwargs["response_model"]) if "response_model" in kwargs else ""
                status = _status_code_to_text(kwargs.get("status_code"))
                full_path = "/api/v1" + (prefix + ("" if route == "/" else route)).replace("//", "/")
                rows.append(
                    {
                        "tag": tag,
                        "method": method,
                        "path": full_path,
                        "status": status,
                        "response": response_model or "-",
                        "description": route_description(tag, method, full_path, fn.name),
                        "security": route_security(tag, full_path),
                    }
                )

    order = {
        "auth": 0,
        "clientes": 1,
        "contactos": 2,
        "oportunidades": 3,
        "cotizaciones": 4,
        "cuentas-cobro": 5,
        "proyectos": 6,
        "proyectos-relaciones": 7,
        "productos": 8,
        "despachos": 9,
        "empleados": 10,
        "paises": 11,
        "fx": 12,
        "notificaciones": 13,
    }
    rows.sort(key=lambda item: (order.get(item["tag"], 99), item["path"], item["method"]))
    return rows


def extract_models() -> list[dict[str, object]]:
    models: list[dict[str, object]] = []
    for file_path in sorted((APP_DIR / "models").glob("*.py")):
        if file_path.name == "__init__.py":
            continue
        tree = ast.parse(file_path.read_text(encoding="utf-8"))
        for cls in [node for node in tree.body if isinstance(node, ast.ClassDef)]:
            table_name = ""
            fields: list[dict[str, str]] = []
            constraints: list[str] = []
            for stmt in cls.body:
                if not isinstance(stmt, ast.Assign):
                    continue

                targets = [target.id for target in stmt.targets if isinstance(target, ast.Name)]
                if not targets:
                    continue
                target = targets[0]

                if target == "__tablename__" and isinstance(stmt.value, ast.Constant):
                    table_name = str(stmt.value.value)
                    continue

                if target == "__table_args__":
                    constraints.append(ast.unparse(stmt.value))
                    continue

                if not isinstance(stmt.value, ast.Call):
                    continue

                call = stmt.value
                func_name = ""
                if isinstance(call.func, ast.Name):
                    func_name = call.func.id
                elif isinstance(call.func, ast.Attribute):
                    func_name = call.func.attr

                if func_name != "Column":
                    continue

                args = [ast.unparse(arg) for arg in call.args]
                column_type = args[0] if args else "Column"
                fk = next((arg for arg in args[1:] if arg.startswith("ForeignKey")), "")
                flags: list[str] = []
                for kw in call.keywords:
                    if kw.arg in {"primary_key", "nullable", "unique", "index", "default", "server_default", "autoincrement"}:
                        flags.append(f"{kw.arg}={ast.unparse(kw.value)}")
                if fk:
                    flags.append(fk)
                fields.append(
                    {
                        "name": target,
                        "type": column_type,
                        "rules": ", ".join(flags) if flags else "-",
                    }
                )

            if table_name:
                models.append(
                    {
                        "table": table_name,
                        "class": cls.name,
                        "file": str(file_path.relative_to(ROOT)),
                        "fields": fields,
                        "constraints": constraints,
                    }
                )
    return models


def extract_schema_contracts() -> list[dict[str, object]]:
    schemas: list[dict[str, object]] = []
    for file_path in sorted((APP_DIR / "schemas").glob("*.py")):
        tree = ast.parse(file_path.read_text(encoding="utf-8"))
        for cls in [node for node in tree.body if isinstance(node, ast.ClassDef)]:
            fields: list[str] = []
            for stmt in cls.body:
                if isinstance(stmt, ast.AnnAssign) and isinstance(stmt.target, ast.Name):
                    default = ""
                    if stmt.value is not None:
                        default = f" = {ast.unparse(stmt.value)}"
                    fields.append(f"{stmt.target.id}: {ast.unparse(stmt.annotation)}{default}")
            if fields:
                schemas.append(
                    {
                        "schema": cls.name,
                        "file": str(file_path.relative_to(ROOT)),
                        "fields": fields,
                    }
                )
    return schemas


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _twips(inches: float) -> int:
    return int(round(inches * 1440))


def set_cell_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(_twips(width)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table, widths: list[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(_twips(sum(widths))))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(_twips(width)))
        grid.append(col)
    tbl.insert(1, grid)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths: list[float]) -> None:
    set_table_fixed_layout(table, widths)
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def set_paragraph_keep(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, name: str | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if name is not None:
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str = ACCENT, size: str = "10", position: str = "bottom") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = borders.find(qn(f"w:{position}"))
    if border is None:
        border = OxmlElement(f"w:{position}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.22)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(10.3)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)
    styles["Normal"].paragraph_format.space_after = Pt(5.5)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 25, ACCENT_DARK),
        ("Heading 1", 14.2, ACCENT_DARK),
        ("Heading 2", 12.2, ACCENT_DARK),
        ("Heading 3", 10.8, TEXT),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(11 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(10)
    styles["List Number"].font.name = "Calibri"
    styles["List Number"].font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_paragraph_keep(p)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def set_cell_text(cell, text: str, font_size: float = 8.6, bold: bool = False, color: str = TEXT, font_name: str = "Aptos") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    set_run_font(run, font_size, bold, color, font_name)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size: float = 8.4, header_fill: str = HEADER_FILL):
    """Render tabular content as stable labeled blocks.

    The artifact renderer used for QA can collapse native Word tables in this
    environment. Labeled blocks keep the same information without layout risk.
    """
    for row_values in rows:
        first = doc.add_paragraph()
        first.paragraph_format.space_before = Pt(2)
        first.paragraph_format.space_after = Pt(1)
        first.paragraph_format.left_indent = Inches(0.12)
        set_paragraph_shading(first, "F8FAFC")

        primary_parts = []
        for idx in range(min(3, len(headers), len(row_values))):
            value = str(row_values[idx]).strip()
            if value:
                primary_parts.append(f"{headers[idx]}: {value}")
        run = first.add_run("  |  ".join(primary_parts))
        set_run_font(run, max(font_size, 7.0), True, ACCENT_DARK)

        details = doc.add_paragraph()
        details.paragraph_format.left_indent = Inches(0.28)
        details.paragraph_format.space_after = Pt(4)
        details.paragraph_format.line_spacing = 1.0
        for idx in range(3, min(len(headers), len(row_values))):
            if idx > 3:
                details.add_run("\n")
            label = details.add_run(f"{headers[idx]}: ")
            set_run_font(label, max(font_size - 0.1, 6.8), True, ACCENT_DARK)
            value = details.add_run(str(row_values[idx]))
            set_run_font(value, max(font_size - 0.1, 6.8), False, TEXT)
    doc.add_paragraph()
    return None


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths: list[float] = [1.8, 4.9]) -> None:
    for key, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(1.5)
        label = p.add_run(f"{key}: ")
        set_run_font(label, 9, True, ACCENT_DARK)
        run = p.add_run(value)
        set_run_font(run, 9, False, TEXT)
    doc.add_paragraph()


def add_note_box(doc: Document, title: str, body: str, fill: str = ACCENT_SOFT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.12)
    set_paragraph_shading(p, fill)
    run = p.add_run(title)
    set_run_font(run, 9.4, True, ACCENT_DARK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.12)
    p2.paragraph_format.space_after = Pt(5)
    set_paragraph_shading(p2, fill)
    r2 = p2.add_run(body)
    set_run_font(r2, 8.8, False, TEXT)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.rstrip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(0)
        set_paragraph_shading(p, "F6F8FA")
        run = p.add_run(line)
        set_run_font(run, 8.2, False, "24292F", "Consolas")
    doc.add_paragraph()


def add_cover_metadata_table(doc: Document) -> None:
    rows = [
        ("Área", "Ingeniería"),
        ("Elaborado por", "Daniel Quiroz"),
        ("Revisado por", "Gerencia"),
        ("Versión", "v2.0"),
        ("Fecha", "12/05/2026"),
        ("Documento base", "INFORME TECNICO API v.1.docx"),
    ]
    for key, value in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1.5)
        label = p.add_run(key)
        set_run_font(label, 9.7, True, TEXT, "Calibri")
        sep = p.add_run("    |    ")
        set_run_font(sep, 9.7, False, "888888", "Calibri")
        content = p.add_run(value)
        set_run_font(content, 9.7, False, TEXT, "Calibri")
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run().add_picture(str(LOGO_PATH), width=Inches(4.1))

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME TÉCNICO / SISTEMAS")
    set_run_font(run, 24, True, ACCENT)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("API v.2")
    set_run_font(run, 20, False, ACCENT_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Documentación técnica actualizada desde la versión 1")
    set_run_font(run, 11.5, False, "666666")

    doc.add_paragraph()
    add_cover_metadata_table(doc)

    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification.add_run("Clasificación: [Confidencial]")
    set_run_font(run, 10.5, False, "666666")

    doc.add_page_break()


def add_static_index(doc: Document) -> None:
    add_heading(doc, "Índice", 1)
    items = [
        "1. Resumen ejecutivo y control de cambios",
        "2. Alcance funcional",
        "3. Arquitectura del backend",
        "4. Tecnologias, dependencias y configuracion",
        "5. Seguridad, autenticacion y permisos",
        "6. Modelo de datos",
        "7. Catalogo de endpoints API v1",
        "8. Reglas de negocio e integraciones",
        "9. Validacion, errores y contratos",
        "10. Operacion local, despliegue y troubleshooting",
        "11. Apendices tecnicos",
    ]
    add_bullets(doc, items)
    doc.add_page_break()


def model_purpose(table: str) -> str:
    purposes = {
        "cliente": "Clientes y datos tributarios/comerciales.",
        "contacto": "Contactos asociados opcionalmente a clientes.",
        "oportunidad": "Oportunidades comerciales y seguimiento previo a cotizacion/proyecto.",
        "cotizacion": "Cotizaciones vigentes, totales y productos serializados.",
        "cotizacion_versiones_v2": "Historico de snapshots previos a actualizaciones de cotizacion.",
        "proyecto": "Proyectos operativos ligados opcionalmente a una oportunidad.",
        "cuenta_cobro": "Datos persistidos para cuentas de cobro.",
        "producto": "Catalogo de productos, costos, imagenes, inventario y metadatos.",
        "pais": "Parametros logisticos por pais para calculos de importacion.",
        "despachos": "Despachos, guias, valores y estados logisticos.",
        "empleado": "Usuarios/empleados, roles inferidos, jefe directo y permisos de vista.",
        "auth_refresh_session": "Sesiones refresh persistidas como hash para rotacion/revocacion.",
        "notificacion": "Notificaciones internas por empleado, entidad y accion requerida.",
        "proyecto_empleado": "Relacion muchos-a-muchos Proyecto-Empleado.",
        "proyecto_cliente": "Relacion muchos-a-muchos Proyecto-Cliente.",
        "proyecto_despacho": "Relacion muchos-a-muchos Proyecto-Despacho.",
    }
    return purposes.get(table, "Tabla del dominio de negocio.")


def relationship_summary() -> list[str]:
    return [
        "Cliente 1:N Oportunidad y Cliente 1:N Contacto.",
        "Oportunidad 1:N Cotizacion y Oportunidad 1:N Proyecto, con regla de unicidad funcional: una oportunidad no debe quedar ligada a mas de un proyecto.",
        "Empleado 1:N Cotizacion; Empleado tambien participa en aprobaciones mediante jefe_id y permisos de vista.",
        "Proyecto M:N Empleado, Cliente y Despacho mediante tablas puente.",
        "Cotizacion 1:N CotizacionVersion; cada actualizacion guarda primero el estado anterior como version historica.",
        "Notificacion apunta a destinatario y actor en empleado, y puede referenciar entidades de negocio por entidad_tipo/entidad_id.",
        "CuentaCobro conserva campos persistidos y expone campos virtuales proyecto_id/oportunidad_id hidratados desde proyecto/cotizacion.",
    ]


def add_summary_sections(doc: Document) -> None:
    add_heading(doc, "1. Resumen Ejecutivo", 1)
    doc.add_paragraph(
        "Este informe actualiza y amplia la primera version de la documentacion tecnica de la API Ingear. "
        "La version actual del backend conserva el nucleo FastAPI + SQLAlchemy, pero incorpora autenticacion JWT, "
        "sesiones refresh, permisos por rol/vista, modulos nuevos de contactos, cuentas de cobro, notificaciones, "
        "tasas FX, paises, versionado de cotizaciones e integraciones con Gmail, Google Drive, World Office, BanRep, BCE y Cloudflare Turnstile."
    )
    add_note_box(
        doc,
        "Correccion importante frente a v1",
        "La advertencia de la version anterior sobre una ForeignKey a world_oficce ya no aplica en el modelo actual. "
        "cliente.world_office_id es un campo String opcional, sin llave foranea declarada.",
        fill="FFF7E6",
    )

    add_heading(doc, "1.1 Control de cambios incorporado", 2)
    changes = [
        ["Autenticacion", "Se documentan login, refresh, logout, perfil y cambio de contrasena con bcrypt, PyJWT y cookie HttpOnly."],
        ["Permisos", "Se agrega matriz de roles y permisos de vista comercial, incluyendo marcador comercial.views.v2."],
        ["Cotizaciones", "Se describe aprobacion por jefe/Gerencia, versionado historico, permisos por equipo, email PDF y creacion de proyecto al pasar a Ganada."],
        ["Productos", "Se documenta calculo de costo_ingear, imagen desde Google Drive y sincronizacion de existencias World Office."],
        ["Cuentas de cobro", "Se agrega prefill desde proyecto, oportunidad, cliente y ultima cotizacion."],
        ["Operaciones", "Se incluyen variables .env actuales, actualizaciones de esquema en startup, errores, troubleshooting y recomendaciones de produccion."],
    ]
    add_table(doc, ["Area", "Cambio agregado o actualizado"], changes, widths=[1.7, 5.5], font_size=8.7)


def add_scope(doc: Document) -> None:
    add_heading(doc, "2. Alcance Funcional", 1)
    doc.add_paragraph("La API soporta el flujo de negocio principal de Ingear:")
    add_bullets(
        doc,
        [
            "Gestion de clientes, contactos y datos tributarios/comerciales.",
            "Gestion de oportunidades comerciales vinculadas a clientes y responsables.",
            "Gestion de cotizaciones, productos cotizados, aprobaciones, historico y envio por correo.",
            "Generacion y administracion de proyectos, estados por area y relaciones con empleados, clientes y despachos.",
            "Gestion de cuentas de cobro con precarga desde proyectos.",
            "Catalogo de productos con costos, PVP, inventario, fichas, imagenes y sincronizacion World Office.",
            "Gestion de despachos, paises logisticos, tasas de cambio y notificaciones internas.",
            "Autenticacion, autorizacion por rol y permisos finos de vistas comerciales.",
        ],
    )


def add_architecture(doc: Document) -> None:
    add_heading(doc, "3. Arquitectura del Backend", 1)
    doc.add_paragraph(
        "El proyecto sigue una arquitectura por capas. FastAPI expone routers versionados; los endpoints aplican dependencias "
        "de autenticacion/autorizacion, delegan persistencia a CRUD/SQLAlchemy y concentran integraciones externas en services."
    )
    add_code_block(
        doc,
        """API_Ingear/
app/
  main.py                  # Crea FastAPI, CORS, router /api/v1 y schema updates
  core/                    # Settings, seguridad JWT/bcrypt, permisos de vista
  db/                      # Engine, SessionLocal, tipos custom y actualizaciones de esquema
  api/v1/api.py            # Enrutador principal
  api/v1/endpoints/        # Endpoints REST por dominio
  models/                  # Modelos SQLAlchemy
  schemas/                 # Contratos Pydantic v2
  crud/                    # Persistencia generica y especializada
  services/                # Integraciones y servicios de dominio"""
    )
    rows = [
        ["Entrada HTTP", "app.main / api.v1.api", "Monta /api/v1, CORS y endpoint raiz /."],
        ["Dependencias", "app.api.deps / app.api.desp", "Obtiene usuario actual y valida roles/permisos."],
        ["Dominio REST", "app.api.v1.endpoints", "Define rutas, reglas de negocio de borde y codigos HTTP."],
        ["Contratos", "app.schemas", "Valida request/response con Pydantic v2 y alias compatibles."],
        ["Persistencia", "app.crud + app.models", "Ejecuta CRUD y reglas transaccionales sobre SQLAlchemy."],
        ["Servicios", "app.services", "Gmail, Turnstile, World Office, tasas FX, imagenes y notificaciones."],
        ["Base de datos", "app.db.session", "Crea engine con pool_pre_ping y fallback psycopg -> pg8000."],
    ]
    add_table(doc, ["Capa", "Modulo", "Responsabilidad"], rows, widths=[1.25, 1.9, 4.05], font_size=8.4)


def add_technology_and_config(doc: Document) -> None:
    add_heading(doc, "4. Tecnologias, Dependencias y Configuracion", 1)
    requirements = parse_requirements()
    req_rows = [[req.split("==")[0].split("[")[0], req] for req in requirements]
    add_heading(doc, "4.1 Stack principal", 2)
    add_table(doc, ["Componente", "Version / paquete"], req_rows, widths=[2.35, 4.85], font_size=8.6)

    add_heading(doc, "4.2 Variables de entorno", 2)
    setting_purpose = {
        "APP_NAME": "Nombre mostrado en OpenAPI y endpoint raiz.",
        "DATABASE_URL": "URL SQLAlchemy. Produccion: PostgreSQL; desarrollo puede usar SQLite si se configura.",
        "GOOGLE_SERVICE_ACCOUNT_FILE": "Ruta del service account para Gmail/Google APIs.",
        "SECRET_KEY": "Clave privada para firmar JWT; debe protegerse fuera del repo.",
        "JWT_ALGORITHM": "Algoritmo de firma JWT.",
        "ACCESS_TOKEN_EXPIRE_MINUTES": "Duracion del access token Bearer.",
        "REFRESH_TOKEN_EXPIRE_HOURS": "Duracion del refresh token persistido.",
        "REFRESH_COOKIE_NAME": "Nombre de la cookie HttpOnly.",
        "REFRESH_COOKIE_PATH": "Path de cookie restringido a auth.",
        "REFRESH_COOKIE_SECURE": "Debe ser true en HTTPS/produccion.",
        "REFRESH_COOKIE_SAMESITE": "Politica SameSite para la cookie.",
        "REFRESH_COOKIE_DOMAIN": "Dominio opcional de cookie.",
        "TURNSTILE_ENABLED": "Activa validacion Cloudflare Turnstile en login.",
        "TURNSTILE_SECRET_KEY": "Secreto de Turnstile.",
        "TURNSTILE_SITEVERIFY_URL": "Endpoint de verificacion Turnstile.",
        "WORLD_OFFICE_ENABLED": "Activa consulta/sincronizacion de inventario World Office.",
        "WORLD_OFFICE_SERVER": "Servidor SQL Server/ODBC de World Office.",
        "WORLD_OFFICE_DATABASE": "Base de datos World Office.",
        "WORLD_OFFICE_USERNAME": "Usuario World Office.",
        "WORLD_OFFICE_PASSWORD": "Clave World Office.",
        "WORLD_OFFICE_ODBC_DRIVER": "Driver ODBC instalado en el host.",
        "WORLD_OFFICE_CONNECTION_TIMEOUT_SECONDS": "Timeout de conexion ODBC.",
        "WORLD_OFFICE_QUERY_TIMEOUT_SECONDS": "Timeout de consultas ODBC.",
        "WORLD_OFFICE_INVENTORY_CACHE_SECONDS": "TTL del cache en memoria de inventario.",
        "WORLD_OFFICE_BODEGA_CODIGO": "Filtro opcional de bodega.",
        "LOGIN_RATE_LIMIT_WINDOW_SECONDS": "Ventana de rate limit de login.",
        "LOGIN_RATE_LIMIT_MAX_ATTEMPTS_PER_IP": "Intentos maximos por IP.",
        "LOGIN_RATE_LIMIT_MAX_ATTEMPTS_PER_USER": "Intentos maximos por usuario.",
        "LOGIN_FAILURE_WINDOW_SECONDS": "Ventana de fallos para bloqueo progresivo.",
        "LOGIN_FAILURE_LOCK_THRESHOLD": "Fallos necesarios para bloquear.",
        "LOGIN_FAILURE_LOCK_BASE_SECONDS": "Primer bloqueo en segundos.",
        "LOGIN_FAILURE_LOCK_BACKOFF_MULTIPLIER": "Multiplicador exponencial del bloqueo.",
        "LOGIN_FAILURE_LOCK_MAX_SECONDS": "Bloqueo maximo en segundos.",
    }
    settings_rows = []
    for field in extract_settings_fields():
        settings_rows.append(
            [
                field["name"],
                field["required"],
                field["default"] or "-",
                setting_purpose.get(field["name"], "Parametro de configuracion del backend."),
            ]
        )
    add_table(doc, ["Variable", "Req.", "Default", "Uso"], settings_rows, widths=[1.75, 0.4, 1.55, 3.5], font_size=7.5)

    add_heading(doc, "4.3 Ejemplo .env seguro", 2)
    add_code_block(
        doc,
        """APP_NAME=Ingear API
DATABASE_URL=postgresql+psycopg://USUARIO:CLAVE@localhost:5432/ingear_db
SECRET_KEY=usar-un-secreto-largo-y-aleatorio
GOOGLE_SERVICE_ACCOUNT_FILE=credentials/service-account.json

REFRESH_COOKIE_SECURE=false
REFRESH_COOKIE_SAMESITE=lax

TURNSTILE_ENABLED=false
WORLD_OFFICE_ENABLED=false"""
    )

    add_heading(doc, "4.4 CORS y arranque", 2)
    add_bullets(
        doc,
        [
            "CORS permite localhost/127.0.0.1 en puertos 5173 y 5174.",
            "Tambien acepta origenes de red local 192.168.x.x, 10.x.x.x y 172.16-31.x.x mediante regex.",
            "En startup se ejecutan actualizaciones idempotentes de esquema para columnas/tablas recientes.",
            "La documentacion OpenAPI queda disponible en /docs y /redoc mientras FastAPI este activo.",
        ],
    )


def add_security(doc: Document) -> None:
    add_heading(doc, "5. Seguridad, Autenticacion y Permisos", 1)
    add_heading(doc, "5.1 Flujo de autenticacion", 2)
    rows = [
        ["Login", "POST /api/v1/auth/login", "Valida rate limit, Turnstile opcional, credenciales bcrypt y estado activo."],
        ["Access token", "JWT Bearer", "Payload con sub=empleado.id y expiracion ACCESS_TOKEN_EXPIRE_MINUTES."],
        ["Refresh", "Cookie HttpOnly", "Token aleatorio almacenado como SHA-256 en auth_refresh_session; se rota en cada refresh."],
        ["Logout", "POST /api/v1/auth/logout", "Revoca sesion activa y elimina cookie."],
        ["Cambio de clave", "POST /api/v1/auth/change-password", "Exige clave actual, minimo 8 caracteres, invalida sesiones y crea una nueva."],
        ["Proteccion login", "LoginProtectionService", "Rate limit por IP/usuario, bloqueo progresivo y header Retry-After en 429."],
    ]
    add_table(doc, ["Elemento", "Implementacion", "Detalle"], rows, widths=[1.25, 1.85, 4.1], font_size=8.3)

    add_heading(doc, "5.2 Roles inferidos", 2)
    add_table(
        doc,
        ["Rol", "Criterio de inferencia"],
        [
            ["GERENCIA", "area/cargo contiene gerencia, gerente o director."],
            ["LOGISTICA", "area/cargo contiene logistica."],
            ["INGENIERIA", "area/cargo contiene ingenieria."],
            ["ADMINISTRACION", "area/cargo contiene administr."],
            ["COMERCIAL", "area/cargo contiene comercial, proyectos o costos."],
            ["OTRO", "No coincide con las reglas anteriores."],
        ],
        widths=[1.45, 5.75],
        font_size=8.5,
    )

    add_heading(doc, "5.3 Permisos de vista comercial", 2)
    add_bullets(
        doc,
        [
            "Gerencia obtiene todos los permisos comerciales por resolucion de permisos.",
            "permisos_vistas se almacena como lista JSON en texto mediante JSONEncodedList.",
            "El marcador comercial.views.v2 evita expandir permisos legados y respeta exactamente la seleccion guardada.",
            "Sin marcador v2, la API expande permisos legados para mantener compatibilidad con usuarios existentes.",
        ],
    )
    add_table(
        doc,
        ["Permiso", "Habilita"],
        [
            ["comercial.cotizador", "Modulo de cotizaciones/cotizador."],
            ["comercial.oportunidades", "Modulo de oportunidades comerciales."],
            ["comercial.proyectos", "Gestion de proyectos comerciales/operativos."],
            ["comercial.cuentas-cobro", "Gestion de cuentas de cobro."],
            ["comercial.clientes", "Gestion de clientes."],
            ["comercial.contactos", "Gestion de contactos."],
            ["comercial.productos", "Catalogo de productos."],
        ],
        widths=[2.1, 5.1],
        font_size=8.5,
    )


def add_data_model(doc: Document, models: list[dict[str, object]]) -> None:
    add_heading(doc, "6. Modelo de Datos", 1)
    doc.add_paragraph(
        "El backend usa SQLAlchemy declarativo. Las tablas principales cubren CRM, cotizaciones, proyectos, productos, autenticacion y notificaciones. "
        "En desarrollo se conserva la creacion/ajuste automatico de ciertos objetos; para produccion se recomienda migrar a Alembic."
    )
    add_heading(doc, "6.1 Relaciones relevantes", 2)
    add_bullets(doc, relationship_summary())

    add_heading(doc, "6.2 Inventario de tablas", 2)
    table_rows = []
    for model in models:
        fields = model["fields"]
        key_fields = []
        for field in fields:
            rules = str(field["rules"])
            if "primary_key=True" in rules or "unique=True" in rules or "ForeignKey" in rules or "index=True" in rules:
                key_fields.append(f"{field['name']} ({rules})")
        constraints = "; ".join(str(c) for c in model["constraints"]) or "-"
        table_rows.append(
            [
                str(model["table"]),
                model_purpose(str(model["table"])),
                "\n".join(key_fields[:8]) or "-",
                constraints,
            ]
        )
    add_table(doc, ["Tabla", "Proposito", "Llaves/indices/FK", "Restricciones"], table_rows, widths=[1.35, 2.25, 2.3, 1.3], font_size=7.3)

    add_heading(doc, "6.3 Diccionario de campos", 2)
    field_rows = []
    for model in models:
        for field in model["fields"]:
            field_rows.append(
                [
                    str(model["table"]),
                    str(field["name"]),
                    str(field["type"]),
                    str(field["rules"]),
                ]
            )
    add_table(doc, ["Tabla", "Campo", "Tipo", "Reglas"], field_rows, widths=[1.35, 1.6, 1.65, 2.6], font_size=6.9)


def add_endpoint_catalog(doc: Document, routes: list[dict[str, str]]) -> None:
    add_heading(doc, "7. Catalogo de Endpoints API v1", 1)
    doc.add_paragraph(
        "Prefijo base: /api/v1. Los endpoints de listado aceptan paginacion skip/limit cuando aplica. "
        "Las respuestas usan JSON salvo la ruta de imagen de producto, que retorna image/jpeg."
    )
    route_rows = [
        [
            route["tag"],
            route["method"],
            route["path"],
            route["description"],
            route["security"],
            f"{route['status']} / {route['response']}",
        ]
        for route in routes
    ]
    add_table(
        doc,
        ["Modulo", "Metodo", "Ruta", "Uso", "Seguridad", "Respuesta"],
        route_rows,
        widths=[0.75, 0.52, 1.7, 2.35, 1.25, 0.63],
        font_size=6.2,
    )


def add_business_rules(doc: Document) -> None:
    add_heading(doc, "8. Reglas de Negocio e Integraciones", 1)
    add_heading(doc, "8.1 Cotizaciones", 2)
    add_bullets(
        doc,
        [
            "Al crear una cotizacion, id_empleado se fuerza al usuario autenticado aunque el contrato acepte alias id_empleado/id_cotizador.",
            "productos se valida contra la tabla producto y se persiste como JSON en texto.",
            "costo_fabrica_override solo puede introducirlo Gerencia; otros usuarios pueden conservar overrides existentes pero no agregar nuevos.",
            "Si el empleado tiene jefe directo activo, la cotizacion nace en estado 1 y genera notificacion de aprobacion; si no, nace en estado 2.",
            "El endpoint aprobar permite aprobar al jefe directo o a Gerencia y resuelve notificaciones pendientes de la entidad.",
            "Cada actualizacion con cambios guarda primero una version historica en cotizacion_versiones_v2.",
            "can_edit se calcula por Gerencia, propietario o pertenencia al mismo equipo comercial anclado por jefe/lider.",
            "Si una cotizacion cambia a etapa Ganada y crear_proyecto_ganada=true, se crea un proyecto para la oportunidad si aun no existe.",
            "El envio por email valida PDF base64, tamano maximo de 20 MB y usa Gmail API con delegacion del correo del empleado.",
        ],
    )

    add_heading(doc, "8.2 Productos y World Office", 2)
    add_bullets(
        doc,
        [
            "codigo_producto es obligatorio y unico; la API devuelve 409 con mensaje especifico cuando se duplica.",
            "Para productos fuera de Colombia, costo_ingear se recalcula como precio_pvp * (1 - descuento_fabricante/100), redondeado a 2 decimales.",
            "GET /productos aplica inventario World Office en memoria si la integracion esta habilitada y configurada.",
            "La sincronizacion World Office consulta Vista_ExistenciasPorBodegas via ODBC, filtra bodega opcional y actualiza cantidad_inventario.",
            "La imagen de producto se genera como thumbnail JPEG desde url_imagen usando el servicio de imagen/Google Drive.",
        ],
    )

    add_heading(doc, "8.3 Proyectos y cuentas de cobro", 2)
    add_bullets(
        doc,
        [
            "Proyecto valida que una oportunidad no quede ligada a mas de un proyecto.",
            "Las relaciones proyecto-empleado, proyecto-cliente y proyecto-despacho son M:N y devuelven 409 si la relacion ya existe.",
            "Cuenta de cobro exige proyecto_id al crear; su prefill busca cliente desde oportunidad o relacion proyecto-cliente y toma la ultima cotizacion de la oportunidad.",
            "Cuenta de cobro persiste campos de cliente/proyecto/cotizacion, mientras proyecto_id y oportunidad_id pueden hidratarse como campos virtuales en respuesta.",
        ],
    )

    add_heading(doc, "8.4 Notificaciones, FX y protecciones externas", 2)
    add_bullets(
        doc,
        [
            "Notificaciones se listan por destinatario autenticado, con filtro solo_no_leidas y limite protegido entre 1 y 200.",
            "FX consulta TRM USD/COP desde BanRep SDMX y USD por EUR desde BCE; EUR/COP se calcula multiplicando ambas tasas y se cachea 30 minutos.",
            "Turnstile puede activarse para login; se omite deliberadamente cuando el host de la solicitud es una IP directa, util para acceso LAN.",
            "Gmail usa service account con delegacion de dominio; si falla, el endpoint responde 502.",
        ],
    )


def add_validation_and_errors(doc: Document) -> None:
    add_heading(doc, "9. Validacion, Errores y Contratos", 1)
    add_heading(doc, "9.1 Codigos HTTP esperados", 2)
    rows = [
        ["200 OK", "Lectura, actualizacion o accion exitosa."],
        ["201 Created", "Creacion de recurso o relacion."],
        ["204 No Content", "Eliminacion o logout sin cuerpo."],
        ["400 Bad Request", "Regla de negocio invalida, relacion inexistente o PDF/captcha invalido."],
        ["401 Unauthorized", "Sin token, token expirado/invalido, credenciales invalidas o sesion expirada."],
        ["403 Forbidden", "Cuenta inactiva, rol/permisos insuficientes o restriccion de aprobacion/override."],
        ["404 Not Found", "Recurso no existe o no pertenece al usuario autenticado."],
        ["409 Conflict", "Unicidad, relacion duplicada o conflicto de versionado."],
        ["422 Unprocessable Entity", "Validacion Pydantic: tipos, campos requeridos o rangos invalidos."],
        ["429 Too Many Requests", "Proteccion de login por rate limit o bloqueo temporal."],
        ["502 Bad Gateway", "Fallo en Gmail, World Office, imagen externa o consulta FX aguas abajo."],
        ["503 Service Unavailable", "Turnstile no disponible o respuesta externa invalida."],
    ]
    add_table(doc, ["Codigo", "Uso"], rows, widths=[1.3, 5.9], font_size=8.5)

    add_heading(doc, "9.2 Convenciones de contrato", 2)
    add_bullets(
        doc,
        [
            "Los modelos Create contienen campos requeridos de creacion; Update usa exclude_unset=True para actualizacion parcial.",
            "Los modelos Out agregan id, fechas calculadas o flags de UI como can_edit/can_duplicate cuando aplica.",
            "CotizacionProductoItem acepta alias de frontend como producto_id, nombreParticion y costoFabricaOverride.",
            "Los montos se modelan como Decimal/Numeric en base de datos y contratos para evitar perdida de precision.",
            "La API no debe documentar secretos reales del .env ni rutas absolutas de credenciales en entregables publicos.",
        ],
    )


def add_operations(doc: Document) -> None:
    add_heading(doc, "10. Operacion Local, Despliegue y Troubleshooting", 1)
    add_heading(doc, "10.1 Ejecucion local", 2)
    add_numbered(
        doc,
        [
            "Crear y activar entorno virtual Python.",
            "Instalar dependencias con python -m pip install -r requirements.txt.",
            "Configurar .env con DATABASE_URL, SECRET_KEY y GOOGLE_SERVICE_ACCOUNT_FILE.",
            "Verificar conexion PostgreSQL o usar una DATABASE_URL compatible para desarrollo.",
            "Ejecutar uvicorn app.main:app --reload --host 127.0.0.1 --port 8000.",
            "Abrir http://127.0.0.1:8000/docs para validar OpenAPI.",
        ],
    )
    add_code_block(
        doc,
        """python -m pip install -r requirements.txt
uvicorn app.main:app --reload --host 127.0.0.1 --port 8000

# Salud
curl http://127.0.0.1:8000/

# Documentacion interactiva
http://127.0.0.1:8000/docs"""
    )

    add_heading(doc, "10.2 Recomendaciones para produccion", 2)
    add_bullets(
        doc,
        [
            "Adoptar Alembic para migraciones y retirar cambios automaticos de esquema fuera de desarrollo.",
            "Usar HTTPS, REFRESH_COOKIE_SECURE=true, SECRET_KEY rotado y almacenamiento seguro de credenciales.",
            "Restringir CORS a dominios reales de frontend; la regex de red local debe limitarse a entornos internos.",
            "Persistir logs estructurados y metricas de errores 4xx/5xx, latencia y fallos de integraciones.",
            "Proteger World Office y Gmail con credenciales de minimo privilegio y rotacion controlada.",
            "Agregar pruebas pytest para CRUD, auth, permisos, cotizaciones, cuentas de cobro e integraciones simuladas.",
            "Definir backups y restauracion para PostgreSQL, especialmente tablas cotizacion, versiones y auth_refresh_session.",
        ],
    )

    add_heading(doc, "10.3 Troubleshooting", 2)
    rows = [
        ["ERR_CONNECTION_REFUSED en /docs", "Uvicorn no esta corriendo o fallo el startup. Revisar uvicorn.stderr.log y variables .env."],
        ["401 Token expirado", "Renovar por /auth/refresh o iniciar sesion nuevamente si la cookie expiro/revoco."],
        ["403 No autorizado", "Verificar area/cargo, permisos_vistas y marcador comercial.views.v2."],
        ["409 producto duplicado", "Cambiar codigo_producto; la restriccion uq_producto_codigo es unica."],
        ["409 relacion duplicada", "La relacion M:N ya existe; eliminar antes de volver a asignar."],
        ["502 World Office", "Validar pyodbc, driver ODBC, servidor, credenciales, bodega y vista Vista_ExistenciasPorBodegas."],
        ["502 Gmail", "Validar service account, delegacion de dominio, correo del empleado y permisos gmail.send."],
        ["422 validacion", "Comparar payload con schemas Pydantic y tipos Decimal/date/int esperados."],
    ]
    add_table(doc, ["Sintoma", "Accion sugerida"], rows, widths=[2.0, 5.2], font_size=8.3)


def add_appendices(doc: Document, schemas: list[dict[str, object]]) -> None:
    add_heading(doc, "11. Apendices Tecnicos", 1)
    add_heading(doc, "11.1 Actualizaciones de esquema en startup", 2)
    add_bullets(
        doc,
        [
            "empleado.permisos_vistas se agrega como TEXT si no existe.",
            "empleado.jefe_id se agrega o se corrige desde jede_id si existe el typo legado.",
            "auth_refresh_session, cuenta_cobro y notificacion se crean con checkfirst=True.",
            "cotizacion_versiones_v2.estado se agrega si la tabla existe y la columna falta.",
        ],
    )

    add_heading(doc, "11.2 Resumen de contratos Pydantic", 2)
    rows = []
    for schema in schemas:
        fields = "\n".join(schema["fields"])
        rows.append([str(schema["schema"]), str(schema["file"]), fields])
    add_table(doc, ["Schema", "Archivo", "Campos"], rows, widths=[1.65, 2.05, 3.5], font_size=6.7)

    add_heading(doc, "11.3 Fuentes revisadas", 2)
    add_bullets(
        doc,
        [
            r"C:\Users\dante\Downloads\INFORME TECNICO API v.1.docx",
            str(ROOT / "app" / "main.py"),
            str(ROOT / "app" / "api" / "v1" / "api.py"),
            str(ROOT / "app" / "api" / "v1" / "endpoints"),
            str(ROOT / "app" / "models"),
            str(ROOT / "app" / "schemas"),
            str(ROOT / "app" / "services"),
            str(ROOT / "requirements.txt"),
        ],
    )


def add_membrete(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        header_p = header.paragraphs[0]
        clear_paragraph(header_p)
        header_p.paragraph_format.space_after = Pt(2)
        header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.1), WD_TAB_ALIGNMENT.RIGHT)
        brand = header_p.add_run("IngeAr")
        set_run_font(brand, 9.4, True, ACCENT, "Calibri")
        header_text = header_p.add_run(" | Informe técnico API")
        set_run_font(header_text, 8.8, False, "666666", "Calibri")
        set_paragraph_border(header_p, ACCENT, size="8", position="bottom")

        footer = section.footer
        footer_p = footer.paragraphs[0]
        clear_paragraph(footer_p)
        footer_p.paragraph_format.space_before = Pt(2)
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_border(footer_p, "B9C7CC", size="4", position="top")
        left = footer_p.add_run("© Ingear - Uso interno")
        set_run_font(left, 8.2, False, "666666", "Calibri")


def build_document() -> None:
    ensure_logo_asset()
    routes = extract_routes()
    models = extract_models()
    schemas = extract_schema_contracts()

    doc = Document()
    style_document(doc)
    doc.core_properties.title = "Informe Tecnico API Ingear v2"
    doc.core_properties.subject = "Documentacion tecnica API"
    doc.core_properties.author = "Ingear"
    doc.core_properties.comments = "Actualizado desde la version v1 con analisis del codigo fuente."
    add_membrete(doc)

    add_cover(doc)
    add_static_index(doc)
    add_summary_sections(doc)
    add_scope(doc)
    add_architecture(doc)
    add_technology_and_config(doc)
    add_security(doc)
    add_data_model(doc, models)
    add_endpoint_catalog(doc, routes)
    add_business_rules(doc)
    add_validation_and_errors(doc)
    add_operations(doc)
    add_appendices(doc, schemas)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(OUTPUT)
    except PermissionError:
        fallback = OUTPUT.with_name("INFORME_TECNICO_API_v2_membrete_IngeAr.docx")
        doc.save(fallback)
        print(fallback)


if __name__ == "__main__":
    build_document()
