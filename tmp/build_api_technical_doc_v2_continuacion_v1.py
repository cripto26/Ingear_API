# -*- coding: utf-8 -*-
from __future__ import annotations

import importlib.util
from collections import defaultdict
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
BASE_SCRIPT = ROOT / "tmp" / "build_api_technical_doc_v2.py"
OUTPUT = ROOT / "docs" / "INFORME_TECNICO_API_v2_actualizado_sobre_v1_IngeAr.docx"
BLUE = "4F81BD"


def load_base_module():
    spec = importlib.util.spec_from_file_location("api_doc_base", BASE_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError("No se pudo cargar el generador base.")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


base = load_base_module()


def find_v1_document() -> Path:
    matches = sorted((Path.home() / "Downloads").glob("INFORME T*API v.1.docx"))
    if not matches:
        raise FileNotFoundError("No se encontró el informe técnico API v.1 en Descargas.")
    return matches[0]


def set_run_font(run, name: str = "Arial", size: float = 12, bold: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), name)


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"No se encontró el párrafo: {text}")


def find_paragraph_starting(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(text):
            return paragraph
    raise ValueError(f"No se encontró el párrafo que inicia con: {text}")


def add_paragraph_like(doc: Document, template, text: str):
    paragraph = doc.add_paragraph(style=template.style)
    paragraph.paragraph_format.space_before = template.paragraph_format.space_before
    paragraph.paragraph_format.space_after = template.paragraph_format.space_after
    paragraph.paragraph_format.line_spacing = template.paragraph_format.line_spacing
    run = paragraph.add_run(text)
    if template.runs:
        src = template.runs[0]
        run.font.name = src.font.name
        run.font.size = src.font.size
        run.font.bold = src.font.bold
        if src.font.color and src.font.color.rgb:
            run.font.color.rgb = src.font.color.rgb
    return paragraph


def _format_like(run, template) -> None:
    if template.runs:
        src = template.runs[0]
        run.font.name = src.font.name
        run.font.size = src.font.size
        run.font.bold = src.font.bold
        if src.font.color and src.font.color.rgb:
            run.font.color.rgb = src.font.color.rgb


def insert_paragraph_like_after(anchor, template, text: str):
    new_p = deepcopy(template._p)
    for child in list(new_p):
        new_p.remove(child)
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    paragraph.style = template.style
    paragraph.paragraph_format.space_before = template.paragraph_format.space_before
    paragraph.paragraph_format.space_after = template.paragraph_format.space_after
    paragraph.paragraph_format.line_spacing = template.paragraph_format.line_spacing
    run = paragraph.add_run(text)
    _format_like(run, template)
    return paragraph


def insert_body_after(anchor, text: str):
    template = find_paragraph_starting(anchor._parent, "Este informe documenta de manera detallada la API del proyecto Ingear.")
    return insert_paragraph_like_after(anchor, template, text)


def insert_list_after(anchor, text: str):
    template = find_paragraph(anchor._parent, "Gestión de clientes (datos tributarios y de contacto).")
    return insert_paragraph_like_after(anchor, template, text)


def insert_heading2_after(anchor, text: str):
    template = next((p for p in anchor._parent.paragraphs if p.style.name == "Heading 2" and p.text.strip()), None)
    if template is None:
        template = find_paragraph(anchor._parent, "1. Alcance funcional (negocio)")
    return insert_paragraph_like_after(anchor, template, text)


def insert_heading1_after(anchor, text: str):
    template = next((p for p in anchor._parent.paragraphs if p.style.name == "Heading 1" and p.text.strip()), None)
    if template is None:
        template = find_paragraph(anchor._parent, "1. Alcance funcional (negocio)")
    return insert_paragraph_like_after(anchor, template, text)


def move_table_after(table, anchor) -> None:
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor._p.addnext(tbl)


def add_section_heading(doc: Document, text: str) -> None:
    template = next((p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.strip()), None)
    if template is None:
        template = find_paragraph(doc, "1. Alcance funcional (negocio)")
    add_paragraph_like(doc, template, text)


def add_subheading(doc: Document, text: str) -> None:
    template = next((p for p in doc.paragraphs if p.style.name == "Heading 2" and p.text.strip()), None)
    if template is None:
        template = find_paragraph(doc, "La API soporta el flujo principal:")
    add_paragraph_like(doc, template, text)


def add_body(doc: Document, text: str) -> None:
    template = find_paragraph_starting(doc, "Este informe documenta de manera detallada la API del proyecto Ingear.")
    add_paragraph_like(doc, template, text)


def add_list_line(doc: Document, text: str) -> None:
    template = find_paragraph(doc, "Gestión de clientes (datos tributarios y de contacto).")
    add_paragraph_like(doc, template, text)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 12) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, "Arial", size, bold)


def set_table_border(table, color: str = BLUE, size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_row_to_existing_table(table, key: str, value: str) -> None:
    row = table.add_row()
    set_cell_text(row.cells[0], key, size=12)
    set_cell_text(row.cells[1], value, size=12)


def add_update_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_border(table)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10.5)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value, size=10)

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)
    for row in table.rows:
        for index, width in enumerate(widths):
            tc_pr = row.cells[index]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def insert_update_table_after(anchor, headers: list[str], rows: list[list[str]], widths: list[int]):
    doc = anchor._parent
    table = doc.add_table(rows=1, cols=len(headers), width=Inches(sum(widths) / 1440))
    set_table_border(table)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10.5)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value, size=10)

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)
    for row in table.rows:
        for index, width in enumerate(widths):
            tc_pr = row.cells[index]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")

    move_table_after(table, anchor)
    return table


def grouped_route_summary(routes: list[dict[str, str]]) -> list[list[str]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for route in routes:
        grouped[route["tag"]].append(route)

    descriptions = {
        "auth": "Login, refresh, logout, usuario actual y cambio de contraseña.",
        "cuentas-cobro": "Gestión y precarga desde proyecto, oportunidad, cliente y cotización.",
        "notificaciones": "Listado, conteo y marcado de notificaciones por usuario.",
        "fx": "Consulta de tasas USD/COP y EUR/COP con BanRep y BCE.",
        "productos": "Catálogo, imagen, costos e inventario World Office.",
        "cotizaciones": "Versionamiento, aprobación, envío por correo y creación de proyecto.",
        "proyectos-relaciones": "Relaciones M:N de proyectos con clientes, despachos y empleados.",
    }

    rows = []
    for tag in sorted(grouped):
        if tag not in descriptions:
            continue
        methods = ", ".join(sorted({route["method"] for route in grouped[tag]}))
        prefixes = sorted({"/".join(route["path"].split("/")[:4]) for route in grouped[tag]})
        rows.append([tag, prefixes[0], methods, descriptions[tag]])
    return rows


def all_route_summary(routes: list[dict[str, str]]) -> list[list[str]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for route in routes:
        grouped[route["tag"]].append(route)
    rows = []
    for tag in sorted(grouped):
        methods = ", ".join(sorted({route["method"] for route in grouped[tag]}))
        prefixes = sorted({"/".join(route["path"].split("/")[:4]) for route in grouped[tag]})
        rows.append([tag, prefixes[0], methods, str(len(grouped[tag]))])
    return rows


def key_settings_rows(settings: list[dict[str, str]]) -> list[list[str]]:
    wanted = [
        "SECRET_KEY",
        "ACCESS_TOKEN_EXPIRE_MINUTES",
        "REFRESH_TOKEN_EXPIRE_HOURS",
        "GOOGLE_SERVICE_ACCOUNT_FILE",
        "TURNSTILE_SECRET_KEY",
        "WORLD_OFFICE_ENABLED",
        "WORLD_OFFICE_ODBC_DRIVER",
        "LOGIN_FAILURE_LOCK_THRESHOLD",
    ]
    by_name = {item["name"]: item for item in settings}
    rows = []
    for name in wanted:
        item = by_name.get(name)
        if not item:
            continue
        default = item.get("default") or "Requerida"
        rows.append([name, item.get("required", ""), default])
    return rows


def new_table_rows() -> list[list[str]]:
    return [
        ["auth_refresh_session", "Sesiones refresh; hash de token, expiración, último uso y revocación."],
        ["contacto", "Contactos asociados a clientes y operaciones comerciales."],
        ["cuenta_cobro", "Cuentas de cobro con precarga desde proyecto/oportunidad/cotización."],
        ["cotizacion_versiones_v2", "Historial de versiones previas a cambios de cotización."],
        ["notificacion", "Notificaciones internas por empleado autenticado."],
        ["pais", "Parámetros logísticos por país para cálculos asociados."],
    ]


def update_cover_and_control(doc: Document) -> None:
    replace_paragraph_text(find_paragraph(doc, "API v.1"), "API v.2")

    metadata = doc.tables[0]
    for row in metadata.rows:
        key = row.cells[0].text.strip()
        if key == "Versión":
            set_cell_text(row.cells[1], "v2.0", size=12)
        elif key == "Fecha":
            set_cell_text(row.cells[1], date.today().strftime("%d/%m/%Y"), size=12)


def update_existing_sections(doc: Document, routes: list[dict[str, str]], settings: list[dict[str, str]]) -> None:
    intro = find_paragraph_starting(doc, "Este informe documenta de manera detallada la API del proyecto Ingear.")
    replace_paragraph_text(
        intro,
        intro.text + " Esta versión continúa el mismo informe y actualiza explícitamente los cambios funcionales y técnicos incorporados en la API v.2.",
    )

    cursor = insert_heading2_after(intro, "Control de cambios incorporado en v2")
    insert_update_table_after(
        cursor,
        ["Área", "Cambio agregado o actualizado"],
        [
            ["Autenticación", "Login, refresh, logout, usuario actual, cambio de contraseña, protección contra fuerza bruta y sesión refresh persistida."],
            ["Permisos", "Autorización por rol/área/cargo y permisos de vista comercial."],
            ["Cotizaciones", "Aprobación por jefe/Gerencia, versionamiento histórico y envío PDF por Gmail."],
            ["Productos", "Imagen de producto, costos e inventario sincronizable desde World Office."],
            ["Cuentas de cobro", "Nuevo módulo con precarga desde proyecto, oportunidad, cliente y cotización."],
            ["Notificaciones", "Nuevo módulo de notificaciones internas por empleado autenticado."],
            ["Tasas FX", "Consulta de TRM USD/COP y cálculo EUR/COP."],
        ],
        [2300, 7400],
    )

    scope_end = find_paragraph(doc, "Gestión de despachos y asignación de despachos a proyectos.")
    cursor = insert_heading2_after(scope_end, "Actualización v2 del alcance funcional")
    for item in [
        "Gestión de contactos como recurso independiente del flujo comercial.",
        "Autenticación de empleados y sesiones seguras con refresh token.",
        "Cuentas de cobro asociadas a proyectos, oportunidades, clientes y cotizaciones.",
        "Notificaciones internas para seguimiento de acciones pendientes.",
        "Tasas de cambio y sincronización de inventario World Office.",
    ]:
        cursor = insert_list_after(cursor, item)

    tech_table = doc.tables[1]
    additions = [
        ("Autenticación", "JWT Bearer, refresh token persistido como hash y cookie HttpOnly."),
        ("Autorización", "Permisos por vista comercial y roles operativos por área/cargo."),
        ("Integración Gmail", "Envío de cotizaciones PDF usando service account y delegación."),
        ("Integración World Office", "Consulta/sincronización de inventario por ODBC."),
        ("Tasas FX", "TRM BanRep y cálculo EUR/COP con Banco Central Europeo."),
        ("Notificaciones", "Notificaciones internas por usuario autenticado."),
    ]
    existing = {row.cells[0].text.strip() for row in tech_table.rows}
    for key, value in additions:
        if key not in existing:
            add_row_to_existing_table(tech_table, key, value)

    notes_end = find_paragraph(doc, "Si cambias de PostgreSQL a SQLite (o viceversa), elimina/crea la DB según corresponda para evitar inconsistencias.")
    cursor = insert_heading2_after(notes_end, "Variables agregadas o relevantes para v2")
    insert_update_table_after(cursor, ["Variable", "Req.", "Valor / uso"], key_settings_rows(settings), [3300, 900, 5500])

    tables_anchor = find_paragraph(doc, "despachos")
    cursor = insert_heading2_after(tables_anchor, "Actualización v2 de tablas principales")
    insert_update_table_after(cursor, ["Tabla", "Uso agregado"], new_table_rows(), [3000, 6600])

    relations_anchor = find_paragraph(doc, "Proyecto M:N Empleado / Cliente / Despacho vía tablas puente.")
    cursor = insert_heading2_after(relations_anchor, "Relaciones adicionales documentadas en v2")
    for item in [
        "Empleado 1:N auth_refresh_session para sesiones refresh revocables.",
        "Empleado 1:N notificacion para avisos internos.",
        "Cotización 1:N cotizacion_versiones_v2 para historial de cambios.",
        "Cuenta de cobro vinculable a proyecto, oportunidad, cliente y cotización.",
    ]:
        cursor = insert_list_after(cursor, item)

    warning_title = find_paragraph(doc, "5.5. Advertencia: ForeignKey a 'world_oficce'")
    replace_paragraph_text(warning_title, "5.5. Actualización v2: integración World Office")
    warning_body = find_paragraph_starting(doc, "El modelo Cliente define una llave foránea world_office_id -> world_oficce.id.")
    replace_paragraph_text(
        warning_body,
        "La advertencia de la versión 1 queda actualizada: el backend actual no depende de una tabla world_oficce para crear el esquema base. La integración con World Office se administra por configuración ODBC y servicios de sincronización de inventario.",
    )

    endpoints_heading = find_paragraph(doc, "7. Endpoints (API v1)")
    replace_paragraph_text(endpoints_heading, "7. Endpoints (API v1 actualizados en v2)")

    endpoint_intro = find_paragraph_starting(doc, "Todas las entidades principales siguen el mismo patrón de rutas:")
    replace_paragraph_text(
        endpoint_intro,
        endpoint_intro.text + " En v2 se agregan módulos no CRUD puros para autenticación, notificaciones, tasas de cambio, sincronización World Office, versionamiento y aprobación de cotizaciones.",
    )
    cursor = insert_heading2_after(endpoint_intro, "7.1.1. Resumen de módulos expuestos en v2")
    insert_update_table_after(cursor, ["Módulo", "Prefijo", "Métodos", "Cant."], all_route_summary(routes), [2200, 3200, 2200, 900])

    response_anchor = find_paragraph(doc, "422 Unprocessable Entity: validación Pydantic fallida (faltan campos requeridos o tipos incorrectos).")
    cursor = insert_heading2_after(response_anchor, "Códigos agregados en v2")
    for item in [
        "401 Unauthorized: access token ausente, inválido o expirado.",
        "403 Forbidden: usuario autenticado sin rol, área, cargo o permiso de vista suficiente.",
        "502 Bad Gateway: falla de integración externa, por ejemplo Gmail, World Office, BanRep o BCE.",
    ]:
        cursor = insert_list_after(cursor, item)

    pydantic_anchor = find_paragraph_starting(doc, "Cada recurso tiene tres esquemas Pydantic:")
    insert_body_after(
        pydantic_anchor,
        "En v2 también se documentan schemas específicos para autenticación, refresh de sesión, notificaciones, tasas FX, cuentas de cobro y versiones de cotización.",
    )

    security = find_paragraph(doc, "Agregar autenticación (JWT / OAuth2) si la API estará expuesta fuera de red interna.")
    replace_paragraph_text(security, "La versión 2 ya incorpora autenticación JWT, refresh token en cookie HttpOnly, logout, cambio de contraseña y control de sesiones.")
    cors = find_paragraph(doc, "Implementar CORS controlado si un frontend web consumirá la API.")
    replace_paragraph_text(cors, "Mantener CORS controlado por entorno, especialmente al pasar de localhost a dominio de producción.")

    tests_anchor = find_paragraph(doc, "Definir linters (ruff/flake8) y formateo (black) para consistencia.")
    cursor = insert_heading2_after(tests_anchor, "9.4. Operación y seguridad v2")
    for item in [
        "Proteger SECRET_KEY, Google service account, Turnstile y World Office fuera del repositorio.",
        "Usar HTTPS, cookie Secure y SameSite adecuado en producción.",
        "Registrar auditoría de aprobación, envío de cotizaciones, cuentas de cobro y cambios de estado.",
        "Monitorear errores 401, 403, 409, 422 y 502 por módulo.",
    ]:
        cursor = insert_list_after(cursor, item)

    trouble_anchor = find_paragraph(doc, "Al asignar un empleado/cliente/despacho a un proyecto, si la relación ya existe se devuelve 409. Solución: no repetir asignación, o eliminar primero la relación y volver a crearla.")
    cursor = insert_heading2_after(trouble_anchor, "11.5. Errores de autenticación y permisos")
    for item in [
        "401: renovar sesión con /api/v1/auth/refresh o iniciar sesión nuevamente.",
        "403: revisar área, cargo, rol inferido, permisos_vistas y marcador comercial.views.v2.",
        "Logout/cambio de contraseña: las sesiones refresh previas pueden quedar revocadas.",
    ]:
        cursor = insert_list_after(cursor, item)
    cursor = insert_heading2_after(cursor, "11.6. Integraciones externas")
    for item in [
        "World Office: validar driver ODBC, servidor, base, usuario, bodega y tiempos de consulta.",
        "Gmail: validar service account, delegación de dominio, correo del empleado y permisos gmail.send.",
        "FX: revisar disponibilidad de BanRep/BCE y caché de tasas.",
    ]:
        cursor = insert_list_after(cursor, item)


def append_v2_continuation(doc: Document, routes: list[dict[str, str]]) -> None:
    doc.add_page_break()
    add_section_heading(doc, "12. Actualización técnica API v.2")
    add_body(
        doc,
        "Esta sección se agrega como continuación directa del informe v1. Mantiene el mismo formato del documento original y consolida los cambios detectados en el backend actual.",
    )

    add_subheading(doc, "12.1. Cambios funcionales incorporados")
    for item in [
        "Autenticación completa con login, refresh, logout, usuario actual y cambio de contraseña.",
        "Control de permisos por vista comercial: cotizador, oportunidades, proyectos, cuentas de cobro, clientes, contactos y productos.",
        "Versionamiento histórico de cotizaciones antes de cada actualización.",
        "Flujo de aprobación de cotizaciones por jefe directo o Gerencia.",
        "Envío de cotizaciones por correo mediante Gmail API.",
        "Precarga de cuentas de cobro desde proyecto, oportunidad, cliente y última cotización.",
        "Notificaciones internas por usuario autenticado.",
        "Consulta de tasas de cambio USD/COP y EUR/COP.",
        "Sincronización de existencias desde World Office.",
    ]:
        add_list_line(doc, item)

    add_subheading(doc, "12.2. Módulos y endpoints agregados o ampliados")
    add_update_table(
        doc,
        ["Módulo", "Prefijo", "Métodos", "Cambio documentado"],
        grouped_route_summary(routes),
        [1700, 2200, 1200, 4300],
    )

    add_subheading(doc, "12.3. Seguridad y sesiones")
    add_body(
        doc,
        "El acceso protegido se realiza con token Bearer. La sesión persistente utiliza refresh token en cookie HttpOnly y una tabla auth_refresh_session con hash del token, fechas de expiración, último uso y revocación.",
    )
    add_body(
        doc,
        "El login puede validar Cloudflare Turnstile y aplica protección contra fuerza bruta mediante límites por usuario/IP, ventana de fallos y bloqueo temporal.",
    )

    add_subheading(doc, "12.4. Cambios en persistencia")
    for item in [
        "Se agregan tablas para auth_refresh_session, cuenta_cobro y notificacion.",
        "Se conserva el modelo principal cliente, contacto, oportunidad, cotizacion, proyecto, producto, despacho y empleado.",
        "Se mantiene el uso de tablas puente proyecto_cliente, proyecto_empleado y proyecto_despacho.",
        "Se agrega historial de cotización en cotizacion_versiones_v2.",
        "Se actualizan columnas de empleado para permisos_vistas y jefe_id cuando el esquema lo requiere.",
    ]:
        add_list_line(doc, item)

    add_subheading(doc, "12.5. Variables de entorno nuevas o relevantes")
    add_update_table(
        doc,
        ["Variable", "Uso"],
        [
            ["SECRET_KEY", "Firma de access tokens JWT."],
            ["GOOGLE_SERVICE_ACCOUNT_FILE", "Ruta de credenciales para Gmail/Google APIs."],
            ["TURNSTILE_SECRET_KEY", "Validación opcional de Turnstile en login."],
            ["WORLD_OFFICE_ENABLED", "Activa la consulta/sincronización de inventario."],
            ["WORLD_OFFICE_ODBC_DRIVER", "Driver ODBC instalado en el servidor."],
            ["LOGIN_FAILURE_LOCK_THRESHOLD", "Número de fallos requeridos para bloquear."],
        ],
        [3200, 6600],
    )

    add_subheading(doc, "12.6. Recomendaciones actualizadas")
    for item in [
        "Usar HTTPS y cookies Secure/SameSite en producción.",
        "Proteger SECRET_KEY, credenciales Google, Turnstile y World Office fuera del repositorio.",
        "Agregar migraciones formales con Alembic para cambios de esquema.",
        "Registrar auditoría de acciones críticas: aprobación, envío de cotización, cambio de estado y cuentas de cobro.",
        "Agregar pruebas de endpoints para autenticación, permisos, cotizaciones, cuentas de cobro e integraciones simuladas.",
    ]:
        add_list_line(doc, item)


def build_document() -> Path:
    routes = base.extract_routes()
    settings = base.extract_settings_fields()

    doc = Document(find_v1_document())
    update_cover_and_control(doc)
    update_existing_sections(doc, routes, settings)
    append_v2_continuation(doc, routes)

    doc.core_properties.title = "Informe Técnico API Ingear v2"
    doc.core_properties.subject = "Continuación del informe técnico API v1"
    doc.core_properties.author = "Ingear"
    doc.core_properties.comments = "Documento v2 construido como continuación del archivo v1."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
