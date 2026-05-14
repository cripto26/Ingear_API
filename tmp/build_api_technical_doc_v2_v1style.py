# -*- coding: utf-8 -*-
from __future__ import annotations

import importlib.util
from collections import defaultdict
from datetime import date
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import tempfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE_SCRIPT = ROOT / "tmp" / "build_api_technical_doc_v2.py"
OUTPUT = ROOT / "docs" / "INFORME_TECNICO_API_v2_estilo_v1_IngeAr.docx"
LOGO_PATH = ROOT / "tmp" / "v1_media" / "image2.png"

TEAL = "004051"
GRAY = "5A5A5A"
BLACK = "000000"
TABLE_BLUE = "4F81BD"
TABLE_GRAY = "B0B0B0"


def load_base_module():
    spec = importlib.util.spec_from_file_location("api_doc_base", BASE_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError("No se pudo cargar el generador base.")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


base = load_base_module()


def find_v1_document() -> Path:
    downloads = Path.home() / "Downloads"
    matches = sorted(downloads.glob("INFORME T*API v.1.docx"))
    if not matches:
        raise FileNotFoundError("No se encontro el informe v1 en Descargas.")
    return matches[0]


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def normal_style(doc: Document):
    for name in ("normal", "Normal"):
        try:
            return doc.styles[name]
        except KeyError:
            continue
    return doc.styles[0]


def set_run_font(run, name: str = "Arial", size: float = 12, bold: bool | None = None, color: str = BLACK) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = rgb(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), name)


def paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = "Arial",
    size: float = 12,
    bold: bool | None = None,
    color: str = BLACK,
    align=None,
    before: float | None = None,
    after: float | None = None,
    line_spacing: float | None = None,
):
    p = doc.add_paragraph()
    p.style = normal_style(doc)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_run_font(run, font, size, bold, color)
    return p


def add_blank(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        paragraph(doc)


def section_title(doc: Document, text: str) -> None:
    paragraph(doc, text, size=12, bold=True, before=10, after=4)


def subsection_title(doc: Document, text: str) -> None:
    paragraph(doc, text, size=12, bold=True, before=6, after=2)


def body_text(doc: Document, text: str, *, after: float = 6) -> None:
    paragraph(doc, text, size=12, after=after)


def list_lines(doc: Document, items: list[str]) -> None:
    for item in items:
        p = paragraph(doc, item, size=12, after=0)
        p.paragraph_format.left_indent = Inches(0.18)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def setup_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.866)
        section.bottom_margin = Inches(0.7875)
        section.left_margin = Inches(0.9056)
        section.right_margin = Inches(0.7875)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    normal = normal_style(doc)
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), "Arial")


def add_page_number(paragraph_obj) -> None:
    run = paragraph_obj.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def set_paragraph_border(paragraph_obj, color: str, position: str = "bottom", size: str = "6") -> None:
    p_pr = paragraph_obj._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{position}"))
    if border is None:
        border = OxmlElement(f"w:{position}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)


def add_clean_membrete(doc: Document) -> None:
    for section in doc.sections:
        header_p = section.header.paragraphs[0]
        header_p.clear()
        header_p.paragraph_format.space_after = Pt(2)
        run = header_p.add_run("Ingear")
        set_run_font(run, "Calibri", 9, True, TEAL)
        run = header_p.add_run(" | Informe")
        set_run_font(run, "Calibri", 9, False, GRAY)
        set_paragraph_border(header_p, TEAL, "bottom", "6")

        footer_p = section.footer.paragraphs[0]
        footer_p.clear()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(2)
        set_paragraph_border(footer_p, TABLE_GRAY, "top", "4")
        run = footer_p.add_run("© Ingear - Uso interno  |  Página ")
        set_run_font(run, "Calibri", 8.5, False, GRAY)
        add_page_number(footer_p)


def set_table_border(table, color: str, size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                set_cell_width(row.cells[index], width)


def clear_cell(cell) -> None:
    for p in cell.paragraphs:
        p.clear()


def cell_text(cell, text: str, *, bold: bool = False, size: float = 12, color: str = BLACK, align=None) -> None:
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, "Arial", size, bold, color)


def add_report_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    font_size: float = 12,
    header_bold: bool = True,
    border_color: str = TABLE_BLUE,
    align=WD_TABLE_ALIGNMENT.LEFT,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = align
    set_table_border(table, border_color)
    set_table_width(table, widths)
    for index, header in enumerate(headers):
        cell_text(table.rows[0].cells[index], header, bold=header_bold, size=font_size)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell_text(row.cells[index], value, size=font_size)
    paragraph(doc, after=6)
    return table


def add_metadata_table(doc: Document) -> None:
    rows = [
        ("Área", "Ingeniería"),
        ("Elaborado por", "Daniel Quiroz"),
        ("Revisado por", "Gerencia"),
        ("Versión", "v2.0"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Documento base", "INFORME TÉCNICO API v.1.docx"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table, TABLE_GRAY, "6")
    set_table_width(table, [2.4, 4.4])
    for row, (key, value) in zip(table.rows, rows):
        cell_text(row.cells[0], key, bold=True, size=12)
        cell_text(row.cells[1], value, size=12)
    paragraph(doc, after=6)


def add_cover(doc: Document) -> None:
    p_logo = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    if LOGO_PATH.exists():
        run = p_logo.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(3.6))
    p_logo.paragraph_format.space_after = Pt(24)

    paragraph(
        doc,
        "INFORME TÉCNICO / SISTEMAS",
        font="Calibri",
        size=28,
        bold=True,
        color=TEAL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=12,
        line_spacing=1.15,
    )
    paragraph(
        doc,
        "API v.2",
        font="Arial",
        size=18,
        color=TEAL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=24,
        line_spacing=1.15,
    )
    add_metadata_table(doc)
    paragraph(
        doc,
        "Clasificación: [Confidencial]",
        font="Calibri",
        size=9,
        bold=False,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=3,
        after=6,
        line_spacing=1.15,
    )
    page_break(doc)


def model_purpose(table_name: str) -> str:
    mapping = {
        "auth_refresh_session": "Sesiones refresh con hash de token, expiracion y revocacion.",
        "cliente": "Maestro de clientes, datos tributarios y atributos comerciales.",
        "contacto": "Contactos asociados al cliente u oportunidad.",
        "cotizacion": "Cotizaciones comerciales, aprobacion, envio y estado del negocio.",
        "cotizacion_versiones_v2": "Historial de versiones antes de cada cambio de cotizacion.",
        "cuenta_cobro": "Cuentas de cobro vinculadas a proyecto, cliente u oportunidad.",
        "despachos": "Registros logisticos y trazabilidad de despacho.",
        "empleado": "Usuarios internos, area, cargo, jefe directo y permisos.",
        "notificacion": "Notificaciones por usuario autenticado.",
        "oportunidad": "Oportunidades comerciales y enlace con cliente/cotizacion/proyecto.",
        "pais": "Catalogo de paises.",
        "producto": "Catalogo tecnico-comercial, costos, inventario e imagen.",
        "proyecto": "Proyectos derivados de oportunidades ganadas.",
        "proyecto_cliente": "Relacion muchos-a-muchos entre proyectos y clientes.",
        "proyecto_despacho": "Relacion muchos-a-muchos entre proyectos y despachos.",
        "proyecto_empleado": "Relacion muchos-a-muchos entre proyectos y empleados.",
    }
    return mapping.get(table_name, "Tabla funcional del dominio.")


def config_summary(settings: list[dict[str, str]]) -> list[list[str]]:
    important = [
        "DATABASE_URL",
        "SECRET_KEY",
        "ACCESS_TOKEN_EXPIRE_MINUTES",
        "REFRESH_TOKEN_EXPIRE_DAYS",
        "TURNSTILE_SECRET_KEY",
        "GOOGLE_SERVICE_ACCOUNT_FILE",
        "GOOGLE_DELEGATED_USER",
        "WORLD_OFFICE_CONNECTION_STRING",
        "WORLD_OFFICE_ENABLED",
        "LOGIN_MAX_FAILED_ATTEMPTS",
        "LOGIN_LOCKOUT_MINUTES",
    ]
    by_name = {field["name"]: field for field in settings}
    rows = []
    for name in important:
        field = by_name.get(name)
        if not field:
            continue
        default = field.get("default") or "Requerida"
        rows.append([name, field.get("type", ""), field.get("required", ""), default])
    return rows


def add_scope(doc: Document) -> None:
    body_text(
        doc,
        "Este informe documenta la version 2 de la API del proyecto Ingear. La estructura visual, membrete, tipografia, "
        "tablas y colores se mantienen alineados con el informe tecnico API v.1, pero el contenido fue actualizado con "
        "la revision del codigo fuente actual.",
    )
    section_title(doc, "1. Alcance funcional (negocio)")
    body_text(doc, "La API soporta el flujo principal:")
    list_lines(
        doc,
        [
            "Gestion de clientes, contactos y datos tributarios.",
            "Gestion de oportunidades comerciales vinculadas a clientes.",
            "Creacion, seguimiento y cierre de proyectos.",
            "Generacion, aprobacion, versionamiento y envio de cotizaciones.",
            "Catalogo de productos con costos, imagenes e inventario.",
            "Administracion de cuentas de cobro y despachos.",
            "Autenticacion, sesiones refresh, permisos por vista y notificaciones.",
            "Integraciones con Gmail, World Office, BanRep y Banco Central Europeo.",
        ],
    )


def add_technologies(doc: Document) -> None:
    section_title(doc, "2. Tecnologías, librerías y estándares")
    body_text(doc, "Estándares y convenciones implementadas:")
    rows = [
        ["Framework API", "FastAPI sobre ASGI."],
        ["Servidor", "Uvicorn en desarrollo y compatible con despliegue ASGI."],
        ["ORM", "SQLAlchemy 2.x con DeclarativeBase y SessionLocal."],
        ["Validación", "Pydantic v2 con schemas Create, Update y Out."],
        ["Base de datos", "PostgreSQL como destino principal; soporte SQLite para desarrollo rapido."],
        ["Seguridad", "JWT Bearer, refresh token HttpOnly, hashing bcrypt y bloqueo por intentos fallidos."],
        ["Integraciones", "Gmail API, World Office por ODBC, BanRep y BCE para tasas de cambio."],
    ]
    add_report_table(doc, ["Clave", "Descripción"], rows, [2.35, 4.45], font_size=12)


def add_structure(doc: Document) -> None:
    section_title(doc, "3. Estructura del proyecto (modulos)")
    body_text(doc, "Estructura y responsabilidad por capa:")
    rows = [
        ["app/main.py", "Inicializa FastAPI, CORS, rutas v1 y actualizaciones de esquema en startup."],
        ["app/api/v1/api.py", "Consolida routers por modulo bajo el prefijo /api/v1."],
        ["app/api/v1/endpoints", "Define endpoints REST, dependencias de seguridad y reglas HTTP."],
        ["app/crud", "Centraliza operaciones de persistencia y consultas por entidad."],
        ["app/models", "Mapeo SQLAlchemy de tablas, relaciones, indices y restricciones."],
        ["app/schemas", "Contratos Pydantic de entrada y salida."],
        ["app/services", "Integraciones externas y servicios de dominio."],
        ["app/core", "Configuración, seguridad JWT y permisos por vista."],
        ["app/db", "Sesion, base declarativa y migraciones ligeras de startup."],
    ]
    add_report_table(doc, ["Ruta", "Responsabilidad"], rows, [2.3, 4.5], font_size=12)


def add_configuration(doc: Document, settings: list[dict[str, str]]) -> None:
    section_title(doc, "4. Configuración y variables de entorno")
    body_text(
        doc,
        "La configuracion se carga desde Settings en app/core/config.py. Las variables marcadas como requeridas deben "
        "existir en el entorno o en el archivo .env antes de iniciar la API.",
    )
    add_report_table(
        doc,
        ["Variable", "Tipo", "Req.", "Valor por defecto / uso"],
        config_summary(settings),
        [2.2, 1.0, 0.65, 3.0],
        font_size=10.5,
    )
    body_text(
        doc,
        "Adicionalmente se parametrizan CORS, cookies seguras, SameSite, algoritmo JWT, credenciales de integraciones y "
        "umbrales de bloqueo por login.",
    )


def add_data_model(doc: Document, models: list[dict[str, object]]) -> None:
    section_title(doc, "5. Base de datos y modelo de persistencia")
    body_text(
        doc,
        "El modelo de datos mantiene entidades comerciales, logisticas, de autenticacion y tablas puente para relaciones "
        "muchos-a-muchos. Las tablas principales identificadas son:",
    )
    rows = []
    for model in sorted(models, key=lambda item: str(item["table"])):
        fields = ", ".join(field["name"] for field in model["fields"][:8])
        if len(model["fields"]) > 8:
            fields += ", ..."
        rows.append([str(model["table"]), str(model["class"]), model_purpose(str(model["table"])), fields])
    add_report_table(
        doc,
        ["Tabla", "Clase", "Uso", "Campos clave"],
        rows,
        [1.55, 1.55, 2.2, 1.5],
        font_size=9.2,
    )
    body_text(
        doc,
        "Relaciones destacadas: cliente-contacto, cliente-oportunidad, oportunidad-cotización, oportunidad-proyecto, "
        "cotización-versiones, proyecto-cliente, proyecto-empleado y proyecto-despacho.",
    )


def add_endpoint_catalog(doc: Document, routes: list[dict[str, str]]) -> None:
    section_title(doc, "6. Endpoints REST documentados")
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for route in routes:
        grouped[route["tag"]].append(route)

    rows = []
    for tag in sorted(grouped):
        methods = sorted({route["method"] for route in grouped[tag]})
        prefixes = sorted({"/".join(route["path"].split("/")[:4]) for route in grouped[tag]})
        rows.append([tag, ", ".join(methods), str(len(grouped[tag])), ", ".join(prefixes)])
    add_report_table(
        doc,
        ["Módulo", "Métodos", "Cant.", "Prefijo principal"],
        rows,
        [1.7, 1.25, 0.7, 3.2],
        font_size=10.5,
    )

    subsection_title(doc, "6.1 Catalogo de rutas")
    rows = []
    for route in sorted(routes, key=lambda item: (item["tag"], item["path"], item["method"])):
        rows.append(
            [
                route["method"],
                route["path"],
                route["description"],
                route["security"],
            ]
        )
    add_report_table(
        doc,
        ["Método", "Ruta", "Descripción", "Seguridad"],
        rows,
        [0.75, 2.3, 2.25, 1.5],
        font_size=8.4,
    )


def add_security(doc: Document) -> None:
    section_title(doc, "7. Seguridad, autenticación y permisos")
    body_text(
        doc,
        "La API utiliza access tokens JWT para autorizacion y refresh tokens persistidos como hash en base de datos. "
        "El refresh se transporta mediante cookie HttpOnly y se revoca en logout, cambio de clave o expiracion.",
    )
    rows = [
        ["Login", "POST /api/v1/auth/login valida credenciales, Turnstile opcional y politicas anti fuerza bruta."],
        ["Refresh", "POST /api/v1/auth/refresh rota la sesion y emite un nuevo access token."],
        ["Logout", "POST /api/v1/auth/logout revoca la sesion vigente y limpia la cookie."],
        ["Usuario actual", "GET /api/v1/auth/me entrega empleado, rol inferido y permisos de vista."],
        ["Permisos", "La autorizacion combina rol/cargo/area con permisos_vistas y marcador comercial.views.v2."],
    ]
    add_report_table(doc, ["Componente", "Comportamiento"], rows, [2.0, 4.8], font_size=12)
    body_text(
        doc,
        "Permisos identificados: comercial.cotizador, comercial.oportunidades, comercial.proyectos, "
        "comercial.cuentas-cobro, comercial.clientes, comercial.contactos y comercial.productos.",
    )


def add_business_rules(doc: Document) -> None:
    section_title(doc, "8. Reglas de negocio e integraciones")
    rows = [
        ["Cotizaciones", "El propietario es el usuario autenticado; se versiona antes de actualizar; la aprobacion requiere jefe directo o Gerencia."],
        ["Productos", "codigo_producto es unico; se calculan costos y se sincronizan existencias desde World Office cuando esta habilitado."],
        ["Oportunidades", "Pueden convertirse en proyecto cuando el estado comercial lo permite."],
        ["Proyectos", "Concentran relaciones con empleados, clientes y despachos mediante tablas puente."],
        ["Cuentas de cobro", "Se pueden prefijar con datos de proyecto, oportunidad, cliente y ultima cotizacion relacionada."],
        ["Gmail", "El envio de cotizacion PDF usa cuenta de servicio y delegacion del empleado autenticado."],
        ["FX", "Consulta TRM USD/COP y calcula EUR/COP con BanRep y Banco Central Europeo."],
        ["Notificaciones", "Se listan, cuentan y marcan como leidas por usuario autenticado."],
    ]
    add_report_table(doc, ["Módulo", "Regla / integración"], rows, [1.65, 5.15], font_size=12)


def add_errors(doc: Document) -> None:
    section_title(doc, "9. Validación, errores y troubleshooting")
    body_text(
        doc,
        "Los payloads se validan con Pydantic. Las respuestas HTTP mantienen codigos semanticos para autenticacion, "
        "autorizacion, duplicidad, integraciones y validacion de datos.",
    )
    rows = [
        ["401", "Token ausente, expirado o invalido. Renovar por refresh o iniciar sesion."],
        ["403", "Usuario autenticado sin rol, cargo, area o permiso suficiente."],
        ["409", "Conflicto por codigo unico, relacion duplicada o regla de negocio."],
        ["422", "Payload no cumple tipo, obligatoriedad o formato esperado por schema."],
        ["502", "Fallo de integracion externa como Gmail, World Office, BanRep o BCE."],
    ]
    add_report_table(doc, ["Código", "Causa y acción"], rows, [1.1, 5.7], font_size=12)


def add_changes(doc: Document) -> None:
    section_title(doc, "10. Cambios incorporados frente a v1")
    list_lines(
        doc,
        [
            "Se agrego autenticacion completa con JWT, refresh token, logout, cambio de clave y usuario actual.",
            "Se documento control de permisos por vista y roles operativos.",
            "Se incorporaron cuentas de cobro, notificaciones y relaciones de proyecto.",
            "Se agrego versionamiento de cotizaciones y flujo de aprobacion.",
            "Se documentaron integraciones con Gmail, World Office, BanRep y Banco Central Europeo.",
            "Se incluyo catalogo de rutas actual, modelo de datos y variables de entorno relevantes.",
            "Se agregaron recomendaciones operativas para despliegue, monitoreo y respaldo.",
        ],
    )


def add_recommendations(doc: Document) -> None:
    section_title(doc, "11. Recomendaciones")
    rows = [
        ["Produccion", "Ejecutar bajo HTTPS, restringir CORS y activar cookies Secure/SameSite segun dominio final."],
        ["Base de datos", "Administrar migraciones formales y backups periodicos de PostgreSQL."],
        ["Secretos", "Guardar JWT, Gmail, Turnstile y World Office en un gestor de secretos."],
        ["Observabilidad", "Persistir logs estructurados, latencia, errores 4xx/5xx y fallos de integracion."],
        ["Pruebas", "Agregar pruebas de CRUD, auth, permisos, cotizaciones, cuentas de cobro e integraciones simuladas."],
    ]
    add_report_table(doc, ["Área", "Recomendación"], rows, [1.5, 5.3], font_size=12)


def add_sources(doc: Document) -> None:
    section_title(doc, "12. Fuentes revisadas")
    list_lines(
        doc,
        [
            str(find_v1_document()),
            str(ROOT / "app" / "main.py"),
            str(ROOT / "app" / "api" / "v1" / "api.py"),
            str(ROOT / "app" / "api" / "v1" / "endpoints"),
            str(ROOT / "app" / "models"),
            str(ROOT / "app" / "schemas"),
            str(ROOT / "app" / "services"),
            str(ROOT / "app" / "core" / "config.py"),
            str(ROOT / "requirements.txt"),
        ],
    )


def sanitize_duplicate_style_ids(path: Path) -> None:
    """Remove duplicate style ids inherited from the v1 template for deterministic rendering."""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    style_id_attr = f"{{{ns['w']}}}styleId"

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=str(path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with ZipFile(path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/styles.xml":
                    root = ET.fromstring(data)
                    seen: set[str] = set()
                    for style in list(root.findall("w:style", ns)):
                        style_id = style.get(style_id_attr)
                        if not style_id:
                            continue
                        if style_id in seen:
                            root.remove(style)
                        else:
                            seen.add(style_id)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(info, data)
        tmp_path.replace(path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def build_document() -> Path:
    if hasattr(base, "ensure_logo_asset"):
        base.ensure_logo_asset()
    routes = base.extract_routes()
    models = base.extract_models()
    settings = base.extract_settings_fields()

    doc = Document()
    setup_document(doc)
    add_clean_membrete(doc)

    doc.core_properties.title = "Informe Técnico API Ingear v2"
    doc.core_properties.subject = "Documentación técnica API"
    doc.core_properties.author = "Ingear"
    doc.core_properties.comments = "Reescrito sobre el estilo del informe tecnico API v1."

    add_cover(doc)
    add_scope(doc)
    add_technologies(doc)
    add_structure(doc)
    add_configuration(doc, settings)
    add_data_model(doc, models)
    add_endpoint_catalog(doc, routes)
    add_security(doc)
    add_business_rules(doc)
    add_errors(doc)
    add_changes(doc)
    add_recommendations(doc)
    add_sources(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
